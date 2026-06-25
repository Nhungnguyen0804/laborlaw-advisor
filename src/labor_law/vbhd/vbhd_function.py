from src.utils.file_utils import save_txt, save_json
from pathlib import Path
from docx import Document
import re
import json
import os
import unicodedata

def extract_docx_text(docx_path):
    doc = Document(docx_path)
    lines = []
   
    for para in doc.paragraphs:
        text = normalize_text(para.text)
        if text:
            lines.append(text)
    xml = doc._element.xml

    match = re.search(
        r'([a-zà-ỹ\s,]+ngày\s+\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4})',
        xml,
        re.IGNORECASE
    )

    if match:
        lines.insert(0, normalize_text(match.group(1)))

    return lines
  
def test_content_doc(docx_files,i):
    doc_file_path = docx_files[i]
    print(f"File thứ {i}: {doc_file_path}")
    lines = extract_docx_text(doc_file_path)
    save_txt(lines,f"src/labor_law/vbhd/text_doc_{i}.txt")
    return doc_file_path

# ****************************************************************
def filter_blld_year_outdate(docx_files):
    filtered_files = []
    for docx_path in docx_files:
        lines = extract_docx_text(docx_path)
        header = "\n".join(lines[:100]).lower()
        can_delete = False 
        patterns = [
            r'bộ luật lao động[^.]{0,80}năm\s+(\d{4})',
            r'căn cứ bộ luật lao động[^.]{0,80}năm\s+(\d{4})'
        ]

        for pattern in patterns:
            matches = re.findall(pattern, header)
            for year_str in matches:
                year = int(year_str)
                if year < 2019:  # BLLĐ 2019 có hiệu lực từ 1/1/2021
                    can_delete = True
                    print(f"Loại: {docx_path} (BLLĐ năm {year})")
                    break
            if can_delete:
                break
        if not can_delete:
            filtered_files.append(docx_path)
      
    return filtered_files



def normalize_text(text):
    if not text:
        return ""

    # Chuẩn hóa unicode
    text = unicodedata.normalize("NFC", text)

    # Chuẩn hóa khoảng trắng
    text = re.sub(r'\s+', ' ', text)

    return text.strip().lower()


HIEU_LUC_PATTERNS = [
    'có hiệu lực thi hành kể từ ngày',
    'có hiệu lực thi hành từ ngày',
    'có hiệu lực kể từ ngày',
    'có hiệu lực từ ngày',
    'hiệu lực thi hành kể từ ngày',
]

DATE_PATTERNS = [
    r'ngày\s+(\d{1,2})\s+tháng\s+(\d{1,2})\s+năm\s+(\d{4})',  # ngày 10 tháng 12 năm 2016
    r'ngày (\d{1,2})/(\d{1,2})/(\d{4})',              # ngày 10/12/2016
    r'ngày (\d{1,2})-(\d{1,2})-(\d{4})',              # ngày 10-12-2016
]


def extract_begin_date(lines):
    # hieu luc cu the 
    # hieu luc tu ngay ky 

    for line in lines:
        has_hieu_luc = any(p in line for p in HIEU_LUC_PATTERNS)
        if not has_hieu_luc:
            continue
        for pattern in DATE_PATTERNS:
            match = re.search(pattern, line)
            if match:
                day, month, year = match.group(1), match.group(2), match.group(3)
                return f"{year}-{month.zfill(2)}-{day.zfill(2)}"
    if is_effective_from_sign_date(lines):
        return extract_sign_date(lines)

    return None






def debug_null_files(docx_files):
    null_files = []
    for path in docx_files:
        try:
            lines = extract_docx_text(path)
            date = extract_begin_date(lines)
            if date is None:
                null_files.append(path)
        except Exception as e:
            print(f"ERROR --> {path.name}: {e}")
 
    total = len(docx_files)
    print(f"\nKết quả: {len(null_files)}/{total} file null ({len(null_files)/total*100:.1f}%)\n")
    print("=" * 70)
    for path in null_files:
        try:
            doc = Document(path)
            # Tìm tất cả paragraph có chứa "hiệu lực"
            hieu_luc_lines = [
                p.text.strip()
                for p in doc.paragraphs
                if "hiệu lực" in p.text.lower() and p.text.strip()
            ]
 
            if not hieu_luc_lines:
                # Không có dòng nào chứa "hiệu lực" → file không có trường này
                print(f"KO CÓ 'hiệu lực' --> {path.name}")
            else:
                print(f"NULL - có 'hiệu lực' nhưng ko match --> {path.name}")
                for line in hieu_luc_lines:
                    print(f"  → {repr(line)}")
 
        except Exception as e:
            print(f"ERROR đọc file --> {path.name}: {e}")
 
        print()
 
    print("=" * 70)
    print(f"--> Tổng null: {len(null_files)} file")
    return null_files
 

NGAY_KY_PATTERNS = [
    r'có hiệu lực từ ngày ký',
    r'có hiệu lực kể từ ngày ký',
    r'có hiệu lực thi hành từ ngày ký',
    r'có hiệu lực thi hành kể từ ngày ký',
    r'có hiệu lực kể từ ngày ký ban hành',
    r'có hiệu lực thi hành kể từ ngày ký ban hành',
]

def is_effective_from_sign_date(lines):
    for line in lines:
        if "hiệu lực" in line and "ngày ký" in line:
            return True
        if "hiệu lực" in line and "ký ban hành" in line:
            return True
        if "hiệu lực" in line and "sau khi ký" in line:
            return True
    return False


def extract_sign_date(lines):
    for line in lines[:3]:
        for pattern in DATE_PATTERNS:
            match = re.search(pattern, line)
            if match:
                day, month, year = match.group(1), match.group(2), match.group(3)
                return f"{year}-{month.zfill(2)}-{day.zfill(2)}"
    return None

RAW_DOC_MODIFY_KEYWORDS = [
    "thay thế",
    "hết hiệu lực",
    "không còn hiệu lực",
    "bãi bỏ",
    "hủy bỏ",
]

# regex nhận diện số hiệu văn bản: ví dụ 34/2018/tt-blđtbxh, 31/2015/nđ-cp
VAN_BAN_SO_HIEU_PATTERN = r'\d{1,5}[/-]\d{4}[/-][a-zà-ỹ0-9\-]+'

DOC_TYPE_PATTERN = r'(luật|nghị quyết|nghị định|thông tư|quyết định|pháp lệnh|bộ luật)'
SO_HIEU_PATTERN = r'\d{1,5}[/-]\d{4}[/-][a-zà-ỹ0-9\-]+'

# Bắt full cụm: loại văn bản (+ tên riêng nếu có) + "số" + số hiệu
FULL_DOC_REF_PATTERN = (
    rf'{DOC_TYPE_PATTERN}'           # loại văn bản
    rf'(?:\s+[a-zà-ỹ\s]+?)?'         # tên/mô tả văn bản (không bắt buộc, lazy)
    rf'\s+số\s+({SO_HIEU_PATTERN})'  # "số" + số hiệu -> group 2
)

def extract_doc_references(line):
    """
    Trích các tham chiếu văn bản đầy đủ trong 1 câu.
    Trả về list dict: {"doc_type": ..., "doc_name": ..., "so_hieu": ...}
    """
    results = []
    for match in re.finditer(FULL_DOC_REF_PATTERN, line, re.IGNORECASE):
        doc_type = match.group(1).strip()
        so_hieu = match.group(2).strip()
        full_match = match.group(0).strip()

        # tên văn bản = phần giữa loại văn bản và "số ..."
        doc_name = full_match
        doc_name = re.sub(rf'^{doc_type}\s*', '', doc_name, flags=re.IGNORECASE)
        doc_name = re.sub(rf'\s*số\s+{re.escape(so_hieu)}$', '', doc_name, flags=re.IGNORECASE)
        doc_name = doc_name.strip()

        results.append({
            "doc_type": doc_type,
            "doc_name": doc_name if doc_name else None,
            "so_hieu": so_hieu,
            "raw": full_match,
        })
    return results


def extract_raw_doc_modify(lines):
    """
    Chỉ lấy câu có keyword quan hệ văn bản (thay thế/hết hiệu lực/bãi bỏ...)
    VÀ có số hiệu văn bản
   
    """
    candidates = []
    for line in lines:
        has_keyword = any(keyword in line for keyword in RAW_DOC_MODIFY_KEYWORDS)
        if not has_keyword:
            continue
        if re.search(VAN_BAN_SO_HIEU_PATTERN, line):
            candidates.append(line)
    return candidates

# tìm van ban het han 

PARTIAL_MARKERS = [
    "một số điều",
    "một số khoản",
    "một số điểm",
    "một số cụm từ",
    "một số nội dung",
    "một số quy định",
]

FULL_END_MARKERS = [
    "hết hiệu lực",
    "không còn hiệu lực",
    "hết hiệu lực toàn bộ",
]

EXCEPTION_MARKERS = [
    "trừ trường hợp quy định tại",
    "trừ quy định tại",
]

def classify_doc_modify(line):
    """
    partial--> câu nói về sửa/thay một phần (không set status=end)
    full--> câu nói văn bản B hết hiệu lực toàn bộ (set status=end)
    unknown--> không xác định được
    """
    is_full = any(marker in line for marker in FULL_END_MARKERS)

    # nếu có "hết hiệu lực" + có cụm "trừ trường hợp..." -> vẫn coi là FULL
    has_exception = any(marker in line for marker in EXCEPTION_MARKERS)

    is_partial = any(marker in line for marker in PARTIAL_MARKERS)

    if is_full and (has_exception or not is_partial):
        return "full"
    if is_partial and not is_full:
        return "partial"
    if is_full:
        return "full"
    return "unknown"


def get_replaced_documents(raw_doc_modify):
    """
    Trả về list các văn bản bị thay thế toàn bộ, kèm loại văn bản + tên.
    [
        {"doc_type": "luật", "doc_name": "giao thông đường bộ", "so_hieu": "23/2008/qh12"},
        {"doc_type": "luật", "doc_name": None, "so_hieu": "35/2018/qh14"},
        ...
    ]
    """
    replaced_docs = []
    seen = set()
    for line in raw_doc_modify:
        if classify_doc_modify(line) != "full":
            continue
        refs = extract_doc_references(line)
        for ref in refs:
            if ref["so_hieu"] not in seen:
                seen.add(ref["so_hieu"])
                replaced_docs.append(ref)
    return replaced_docs


SUA_DOI_PATTERN = r'sửa đổi[^.]*?theo\s+(.*?)(?=hết hiệu lực|$)'

def extract_document_relations(line):
    """
    Tách văn bản gốc và (các) văn bản sửa đổi nó.
    Trả về (root_refs, amending_refs)
    """
    m = re.search(SUA_DOI_PATTERN, line, re.IGNORECASE)
    if not m:
        return extract_doc_references(line), []

    amend_part = m.group(1)
    amend_start = line.find(amend_part)

    root_part = line[:amend_start]
    root_refs = extract_doc_references(root_part)
    amending_refs = extract_doc_references(amend_part)
    return root_refs, amending_refs




def build_json(docx_files, output_path,**new_attr):
    json_content = []
    for docx_path in docx_files:
        lines = extract_docx_text(docx_path)
        raw_doc_modify = extract_raw_doc_modify(lines)
        replaced_documents = get_replaced_documents(raw_doc_modify)
        res_replaced_documents = [doc["so_hieu"] for doc in replaced_documents]
        item = {
            "filename": os.path.basename(docx_path),  # chỉ lấy tên file
            "filepath": str(docx_path),
            "begin_date": extract_begin_date(lines),
            "raw_doc_modify": raw_doc_modify,  
            "replaced_documents": res_replaced_documents,   # tất cả VB hết hiệu lực
        }
        
        item.update(new_attr)
        json_content.append(item)
    save_json(json_content,output_path)
    print(f"({len(json_content)} files)")
    return output_path
